from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Paths ───────────────────────────────────────────────────────────────────
LOGO_PATH = r"C:\Users\Administrator\Documents\IAM_BIDDING\Assets\1.png"
OUTPUT    = r"C:\Users\Administrator\Desktop\Linter_User_Manual.docx"

# ── Colors ──────────────────────────────────────────────────────────────────
CYAN      = RGBColor(0x00, 0xD4, 0xE8)
DARK_BG   = RGBColor(0x0D, 0x11, 0x17)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
MID_GREY  = RGBColor(0x44, 0x4A, 0x5A)
LIGHT_BG  = RGBColor(0xF2, 0xFC, 0xFD)
RULE_LINE = RGBColor(0x00, 0xD4, 0xE8)
TABLE_HDR = RGBColor(0x0D, 0x11, 0x17)

doc = Document()

# ── Page Setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(0.75)

# ── Default paragraph style ──────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_TEXT

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   kwargs[side].get('val',   'single'))
            border.set(qn('w:sz'),    kwargs[side].get('sz',    '4'))
            border.set(qn('w:space'), kwargs[side].get('space', '0'))
            border.set(qn('w:color'), kwargs[side].get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(16)
        run.font.color.rgb = DARK_TEXT
        # Cyan underline rule
        p.paragraph_format.space_after = Pt(4)
        # Add bottom border to paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '00D4E8')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run.font.size  = Pt(13)
        run.font.color.rgb = MID_GREY
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = DARK_TEXT
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK_TEXT
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    return p

def add_step_box(number, title, body_lines):
    """Cyan-accented step box."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    w_twips = int(6.5 * 1440)
    tbl.columns[0].width = Emu(int(0.55 * 914400))
    tbl.columns[1].width = Emu(int(5.95 * 914400))

    row = tbl.rows[0]
    row.height = None

    # Number cell
    num_cell = row.cells[0]
    shade_cell(num_cell, '0D1117')
    set_cell_border(num_cell,
        top={'val':'single','sz':'4','color':'00D4E8'},
        left={'val':'single','sz':'4','color':'00D4E8'},
        bottom={'val':'single','sz':'4','color':'00D4E8'},
        right={'val':'nil','sz':'0','color':'FFFFFF'})
    num_p = num_cell.paragraphs[0]
    num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_p.paragraph_format.space_before = Pt(6)
    num_p.paragraph_format.space_after  = Pt(6)
    num_run = num_p.add_run(str(number))
    num_run.bold = True
    num_run.font.size  = Pt(20)
    num_run.font.color.rgb = CYAN

    # Content cell
    body_cell = row.cells[1]
    shade_cell(body_cell, 'F2FCFD')
    set_cell_border(body_cell,
        top={'val':'single','sz':'4','color':'00D4E8'},
        left={'val':'nil','sz':'0','color':'FFFFFF'},
        bottom={'val':'single','sz':'4','color':'00D4E8'},
        right={'val':'single','sz':'4','color':'00D4E8'})
    body_cell.paragraphs[0].clear()
    title_p = body_cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(2)
    title_p.paragraph_format.left_indent  = Inches(0.1)
    t_run = title_p.add_run(title)
    t_run.bold = True
    t_run.font.size  = Pt(12)
    t_run.font.color.rgb = DARK_TEXT

    for line in body_lines:
        bp = body_cell.add_paragraph(line)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after  = Pt(3)
        bp.paragraph_format.left_indent  = Inches(0.1)
        for run in bp.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = MID_GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_two_col_table(headers, rows_data, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths is None:
        each = int(6.5 * 1440 / n_cols)
        col_widths = [each] * n_cols

    for i, cell in enumerate(tbl.rows[0].cells):
        shade_cell(cell, '0D1117')
        set_cell_border(cell,
            top={'val':'single','sz':'4','color':'00D4E8'},
            left={'val':'single','sz':'4','color':'444A5A'},
            bottom={'val':'single','sz':'4','color':'00D4E8'},
            right={'val':'single','sz':'4','color':'444A5A'})
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.05)
        run = p.add_run(headers[i])
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = CYAN

    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx + 1]
        fill = 'FFFFFF' if r_idx % 2 == 0 else 'F7FEFF'
        for c_idx, cell in enumerate(row.cells):
            shade_cell(cell, fill)
            set_cell_border(cell,
                top={'val':'single','sz':'2','color':'CCEEEE'},
                left={'val':'single','sz':'2','color':'CCEEEE'},
                bottom={'val':'single','sz':'2','color':'CCEEEE'},
                right={'val':'single','sz':'2','color':'CCEEEE'})
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.05)
            run = p.add_run(row_data[c_idx])
            run.font.size  = Pt(10.5)
            run.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)
    # Light box background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF8FA')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x70)
    return p

def add_spacer(pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)


# ════════════════════════════════════════════════════════════════════════════
#  HEADER BANNER
# ════════════════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

logo_cell    = tbl.rows[0].cells[0]
title_cell   = tbl.rows[0].cells[1]

shade_cell(logo_cell,  '0D1117')
shade_cell(title_cell, '0D1117')
for c in [logo_cell, title_cell]:
    set_cell_border(c,
        top={'val':'nil'}, left={'val':'nil'},
        bottom={'val':'single','sz':'8','color':'00D4E8'},
        right={'val':'nil'})

# Logo image
logo_p = logo_cell.paragraphs[0]
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.paragraph_format.space_before = Pt(6)
logo_p.paragraph_format.space_after  = Pt(6)
logo_p.paragraph_format.left_indent  = Inches(0.1)
if os.path.exists(LOGO_PATH):
    logo_p.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
else:
    r = logo_p.add_run("RAYBURN IDEAS")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = CYAN

# Title text
title_p = title_cell.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
title_p.paragraph_format.space_before = Pt(14)
title_p.paragraph_format.space_after  = Pt(2)
title_p.paragraph_format.right_indent = Inches(0.1)
t1 = title_p.add_run("Writing Constraint Linter")
t1.bold = True; t1.font.size = Pt(16); t1.font.color.rgb = WHITE
sub_p = title_cell.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sub_p.paragraph_format.right_indent = Inches(0.1)
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(10)
t2 = sub_p.add_run("User Manual  |  Rayburn Ideas")
t2.font.size = Pt(9); t2.font.color.rgb = CYAN

add_spacer(10)

# ════════════════════════════════════════════════════════════════════════════
#  WHAT THIS TOOL DOES
# ════════════════════════════════════════════════════════════════════════════
add_heading("What This Tool Does")
add_body(
    "The Writing Constraint Linter checks any text against seven writing rules. "
    "It flags violations by category. Use it before sending any bid, email, "
    "LinkedIn post, or document."
)

# URL highlight box
url_tbl = doc.add_table(rows=1, cols=1)
url_tbl.style = 'Table Grid'
url_cell = url_tbl.rows[0].cells[0]
shade_cell(url_cell, 'E8FAFB')
set_cell_border(url_cell,
    top={'val':'single','sz':'4','color':'00D4E8'},
    left={'val':'single','sz':'12','color':'00D4E8'},
    bottom={'val':'single','sz':'4','color':'00D4E8'},
    right={'val':'single','sz':'4','color':'CCEEEE'})
up = url_cell.paragraphs[0]
up.paragraph_format.space_before = Pt(6)
up.paragraph_format.space_after  = Pt(2)
up.paragraph_format.left_indent  = Inches(0.1)
ur1 = up.add_run("Live URL:  ")
ur1.bold = True; ur1.font.size = Pt(10.5); ur1.font.color.rgb = MID_GREY
ur2 = up.add_run("https://ericrayburn.github.io/iam-tools/linter.html")
ur2.font.size = Pt(10.5); ur2.font.color.rgb = CYAN; ur2.bold = True
up2 = url_cell.add_paragraph()
up2.paragraph_format.space_before = Pt(0)
up2.paragraph_format.space_after  = Pt(6)
up2.paragraph_format.left_indent  = Inches(0.1)
ur3 = up2.add_run("No login. No install. Works on any browser, any computer.")
ur3.italic = True; ur3.font.size = Pt(10); ur3.font.color.rgb = MID_GREY
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════════════════════
#  STEP BY STEP
# ════════════════════════════════════════════════════════════════════════════
add_heading("Step by Step")
add_spacer(4)

add_step_box(1, "Open the linter",
    ["Go to the URL above in any browser.",
     "You will see two panels. Left panel is for your text. Right panel shows results."])

add_step_box(2, "Paste your text",
    ["Click inside the left panel.",
     "Paste the text you want to check.",
     "The word counter at the bottom updates as you type."])

add_step_box(3, "Run the check",
    ["Press Ctrl + Enter on your keyboard.",
     "Or click the Run Check button.",
     "Results appear instantly in the right panel."])

add_step_box(4, "Read the score",
    ["The top of the right panel shows a score circle.",
     "Green check = no violations. Clear to send.",
     "Orange number = 1 to 4 violations. Review before sending.",
     "Red number = 5 or more violations. Do not send."])

add_step_box(5, "Review each violation group",
    ["Seven rule categories appear below the score.",
     "Red badge = violations found. The group opens automatically.",
     "Green check = that rule passed. The group stays collapsed.",
     "Click any group header to open or close it."])

add_step_box(6, "Read each violation",
    ["Each violation shows two lines.",
     "Top line (orange text): the exact word or phrase that triggered the rule.",
     "Bottom line (grey text): the surrounding sentence for context."])

add_step_box(7, "Fix and re-run",
    ["Go back to the left panel. Edit your text.",
     "Press Ctrl + Enter to re-run the check.",
     "Repeat until the score circle turns green."])

add_step_box(8, "Clear and start over",
    ["Click the Clear button to reset both panels for new text."])


# ════════════════════════════════════════════════════════════════════════════
#  SCORE REFERENCE
# ════════════════════════════════════════════════════════════════════════════
add_heading("Score Reference")
add_two_col_table(
    ["Score Circle", "What It Means", "Action"],
    [
        ["Green  ✓",  "Zero violations",               "Clear to send"],
        ["Orange  1-4", "Minor violations found",       "Review each flag, fix, and re-run"],
        ["Red  5+",   "Multiple violations found",      "Do not send. Fix all flags first"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
#  SEVEN RULES
# ════════════════════════════════════════════════════════════════════════════
add_heading("The Seven Rules")
add_two_col_table(
    ["Rule", "What Triggers It"],
    [
        ["Banned Words",
         "Any of the 38 banned terms: passionate, empower, excited, solution, leverage, synergy, paradigm, robust, seamless, and 29 more"],
        ["Contractions",
         "do not -> don't, we will -> we'll, it is -> it's, cannot -> can't, and all others"],
        ["Em / En Dashes",
         "The long dash character (—), the medium dash (–), or double hyphens (--)"],
        ["Sentences Over 20 Words",
         "Any sentence with more than 20 words. Split it into two."],
        ["Hyphens in Compound Modifiers",
         "Any word-hyphen-word pattern. Example: 600-user, Zero-Touch, PE-backed"],
        ["Passive Voice",
         "is managed, was built, are handled, is known, was given, and similar patterns"],
        ["Hedging Language",
         "hopefully, ideally, may help, would be, might help, and similar soft phrases"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
#  HYPHEN RULE NOTE
# ════════════════════════════════════════════════════════════════════════════
add_heading("Note on the Hyphen Rule")
add_body(
    "The hyphens check flags every hyphenated word pair. Not all hyphens are violations. "
    "Review each one manually."
)

hbl = doc.add_paragraph()
hbl.paragraph_format.space_before = Pt(2)
hbl.paragraph_format.space_after  = Pt(2)
hbl.paragraph_format.left_indent  = Inches(0.3)
r1 = hbl.add_run("Remove: ")
r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)
r2 = hbl.add_run("If the hyphenated words modify a noun that follows. Example:")
r2.font.size = Pt(11); r2.font.color.rgb = DARK_TEXT
add_code_block('  "600-user environment"  ->  "600 user environment"')

hbl2 = doc.add_paragraph()
hbl2.paragraph_format.space_before = Pt(2)
hbl2.paragraph_format.space_after  = Pt(2)
hbl2.paragraph_format.left_indent  = Inches(0.3)
r3 = hbl2.add_run("Keep: ")
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x00, 0x80, 0x40)
r4 = hbl2.add_run("Proper nouns, version numbers, or words not modifying a noun.")
r4.font.size = Pt(11); r4.font.color.rgb = DARK_TEXT
add_code_block('  "PowerShell 7.4"  ->  no hyphen to flag')
add_spacer(4)


# ════════════════════════════════════════════════════════════════════════════
#  UPDATING THE TOOL
# ════════════════════════════════════════════════════════════════════════════
add_heading("Updating the Tool")
add_body("If writing constraints change, edit the local source file and push to GitHub.")
add_body("Local file:")
add_code_block("  C:\\Users\\Administrator\\Documents\\iam-tools\\linter.html")
add_body("Push command (run in PowerShell):")
add_code_block(
    "  cd C:\\Users\\Administrator\\Documents\\iam-tools\n"
    "  git add linter.html\n"
    "  git commit -m \"Update writing constraints\"\n"
    "  git push"
)
add_body("The live URL updates within 2 minutes of the push.")

add_spacer(10)

# ════════════════════════════════════════════════════════════════════════════
#  FOOTER LINE
# ════════════════════════════════════════════════════════════════════════════
fp = doc.add_paragraph()
pPr = fp._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_bdr = OxmlElement('w:top')
top_bdr.set(qn('w:val'),   'single')
top_bdr.set(qn('w:sz'),    '4')
top_bdr.set(qn('w:space'), '4')
top_bdr.set(qn('w:color'), '00D4E8')
pBdr.append(top_bdr)
pPr.append(pBdr)
fp.paragraph_format.space_before = Pt(10)
fp.paragraph_format.space_after  = Pt(4)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Rayburn Ideas  |  Identity and Access Management  |  ericrayburn.apply@gmail.com")
fr.font.size = Pt(9); fr.font.color.rgb = MID_GREY

# ════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
